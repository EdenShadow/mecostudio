#!/usr/bin/env python3
"""
DOCX to Markdown converter
Usage: python convert_docx.py <input.docx> [output.md]
"""

import sys
import os
from pathlib import Path

def convert_docx_to_md(input_path, output_path=None):
    """Convert DOCX file to Markdown."""
    
    try:
        from docx import Document
    except ImportError:
        print("ERROR: python-docx not installed. Run: pip install python-docx")
        sys.exit(1)
    
    input_path = Path(input_path)
    if not input_path.exists():
        print(f"ERROR: File not found: {input_path}")
        sys.exit(1)
    
    # Default output path
    if output_path is None:
        output_path = input_path.with_suffix('.md')
    else:
        output_path = Path(output_path)
    
    # Ensure output directory exists
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Load document
    doc = Document(str(input_path))
    
    md_lines = []
    
    # Add title from filename if no clear title in doc
    title = input_path.stem
    md_lines.append(f"# {title}\n")
    md_lines.append(f"\n*Source: {input_path.name}*\n")
    md_lines.append("---\n")
    
    # Process paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            md_lines.append("\n")
            continue
        
        # Check if it's a heading
        if para.style and para.style.name and para.style.name.startswith('Heading'):
            try:
                level = int(para.style.name.replace('Heading ', ''))
                md_lines.append(f"{'#' * level} {text}\n")
            except:
                md_lines.append(f"## {text}\n")
        else:
            md_lines.append(f"{text}\n")
    
    # Process tables
    for table in doc.tables:
        md_lines.append("\n")
        for i, row in enumerate(table.rows):
            cells = [cell.text.strip() for cell in row.cells]
            md_lines.append("| " + " | ".join(cells) + " |\n")
            if i == 0:  # Header separator
                md_lines.append("|" + "|".join([" --- " for _ in cells]) + "|\n")
        md_lines.append("\n")
    
    # Write output
    with open(output_path, 'w', encoding='utf-8') as f:
        f.writelines(md_lines)
    
    print(f"✓ Converted: {input_path} → {output_path}")
    return output_path

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python convert_docx.py <input.docx> [output.md]")
        sys.exit(1)
    
    input_file = sys.argv[1]
    output_file = sys.argv[2] if len(sys.argv) > 2 else None
    convert_docx_to_md(input_file, output_file)
